from docx import Document
import pandas as pd

def convert_word_to_excel(file_path):

    doc = Document(file_path)
    output_path = file_path.replace(".docx", "_converted.xlsx")

    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        start_row = 0

        # ======================
        # PARAGRAPHS
        # ======================
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                pd.DataFrame([[text]]).to_excel(
                    writer,
                    startrow=start_row,
                    index=False,
                    header=False
                )
                start_row += 1

        start_row += 1  # spacing

        # ======================
        # TABLES
        # ======================
        for table in doc.tables:

            rows = []
            max_cols = 0

            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                max_cols = max(max_cols, len(row_data))
                rows.append(row_data)

            # ❗ Fix uneven columns
            for row in rows:
                if len(row) < max_cols:
                    row.extend([""] * (max_cols - len(row)))

            if not rows:
                continue

            header = rows[0]
            data = rows[1:] if len(rows) > 1 else []

            df = pd.DataFrame(data, columns=header)

            # Remove empty rows/columns
            df = df.dropna(axis=1, how='all')
            df = df.dropna(how='all')

            df.to_excel(writer, startrow=start_row, index=False)

            start_row += len(df) + 3  # spacing between tables

    return output_path